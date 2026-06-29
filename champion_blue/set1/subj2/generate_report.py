# -*- coding: utf-8 -*-
"""[연습세트01·블루] 2과목 데이터분석 — 차트 5종 생성 및 분석 보고서(docx) 작성."""
import os
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from sklearn.model_selection import train_test_split
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import RandomForestClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.neighbors import KNeighborsClassifier
from sklearn.metrics import f1_score
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

plt.rcParams["font.family"] = "Malgun Gothic"
plt.rcParams["axes.unicode_minus"] = False

BASE = os.path.dirname(os.path.abspath(__file__))
ATT = os.path.join(BASE, "첨부")
IMG = os.path.join(BASE, "images")
os.makedirs(IMG, exist_ok=True)

# ---------- 데이터 정제 (문제 스펙 그대로) ----------
a = pd.read_csv(os.path.join(ATT, "청렴지수_현황.csv"), encoding="utf-8-sig")
b = pd.read_csv(os.path.join(ATT, "부패신고_현황.csv"), encoding="utf-8-sig")

m = pd.merge(a, b, on="기관코드", how="inner")   # 175
m = m.dropna()                                    # 148

def iqr_mask(s):
    q1, q3 = s.quantile(0.25), s.quantile(0.75)
    iqr = q3 - q1
    return (s >= q1 - 1.5 * iqr) & (s <= q3 + 1.5 * iqr)

mask = iqr_mask(m["설문인원"]) & iqr_mask(m["평균처리일수"])
clean = m[mask].reset_index(drop=True)            # 141

# ---------- 라벨 생성 (시드 251) ----------
rng = np.random.default_rng(251)
risk = (-0.45 * clean["청렴지수"] + 0.08 * clean["신고건수"]
        - 0.012 * clean["처리율"] + 0.04 * clean["평균처리일수"])
risk_n = (risk - risk.mean()) / risk.std()
noise = rng.normal(0, 0.5, len(clean))
clean["부패위험"] = ((risk_n.values + noise) > 0.1).astype(int)

# ---------- 차트 1: 기관유형별 청렴지수 분포 (boxplot) ----------
fig, ax = plt.subplots(figsize=(7, 4.5))
clean.boxplot(column="청렴지수", by="기관유형", ax=ax, grid=False)
ax.set_title("기관유형별 청렴지수 분포")
ax.set_xlabel("기관유형"); ax.set_ylabel("청렴지수")
plt.suptitle("")
fig.savefig(os.path.join(IMG, "chart1_boxplot.png"), bbox_inches="tight", dpi=120)
plt.close(fig)

# ---------- 차트 2: 청렴지수 히스토그램 ----------
fig, ax = plt.subplots(figsize=(7, 4.5))
ax.hist(clean["청렴지수"], bins=20, color="#4C72B0", edgecolor="white")
ax.set_title("청렴지수 히스토그램")
ax.set_xlabel("청렴지수"); ax.set_ylabel("빈도")
fig.savefig(os.path.join(IMG, "chart2_histogram.png"), bbox_inches="tight", dpi=120)
plt.close(fig)

# ---------- 차트 3: 신고건수 vs 처리율 산점도 ----------
fig, ax = plt.subplots(figsize=(7, 4.5))
sc = ax.scatter(clean["신고건수"], clean["처리율"], c=clean["부패위험"],
                cmap="coolwarm", alpha=0.7, edgecolor="k", linewidth=0.3)
ax.set_title("신고건수 vs 처리율 (색: 부패위험)")
ax.set_xlabel("신고건수"); ax.set_ylabel("처리율(%)")
legend = ax.legend(*sc.legend_elements(), title="부패위험")
ax.add_artist(legend)
fig.savefig(os.path.join(IMG, "chart3_scatter.png"), bbox_inches="tight", dpi=120)
plt.close(fig)

# ---------- 모델 학습 (Q5와 동일 분할) ----------
features = ["청렴지수", "응답률", "설문인원", "신고건수", "처리완료건수", "처리율", "평균처리일수"]
X = clean[features]
y = clean["부패위험"]
Xtr, Xte, ytr, yte = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y)

models = {
    "DecisionTree": DecisionTreeClassifier(random_state=42),
    "RandomForest": RandomForestClassifier(random_state=42),
    "LogisticReg": LogisticRegression(max_iter=1000),
    "KNN": KNeighborsClassifier(),
}
f1_scores = {}
for name, mdl in models.items():
    mdl.fit(Xtr, ytr)
    f1_scores[name] = f1_score(yte, mdl.predict(Xte))

# ---------- 차트 4: 4개 알고리즘 F1 비교 (bar) ----------
fig, ax = plt.subplots(figsize=(7, 4.5))
names = list(f1_scores.keys()); vals = list(f1_scores.values())
bars = ax.bar(names, vals, color=["#4C72B0", "#55A868", "#C44E52", "#8172B2"])
ax.set_title("4개 알고리즘 F1 점수 비교")
ax.set_ylabel("F1 score"); ax.set_ylim(0, 1)
for bar, v in zip(bars, vals):
    ax.text(bar.get_x() + bar.get_width() / 2, v + 0.02, f"{v:.3f}", ha="center")
fig.savefig(os.path.join(IMG, "chart4_bar_f1.png"), bbox_inches="tight", dpi=120)
plt.close(fig)

# ---------- 차트 5: DecisionTree max_depth별 F1 변화 (line) ----------
depths = list(range(1, 16))
depth_f1 = []
for d in depths:
    dt = DecisionTreeClassifier(max_depth=d, random_state=42)
    dt.fit(Xtr, ytr)
    depth_f1.append(f1_score(yte, dt.predict(Xte)))
fig, ax = plt.subplots(figsize=(7, 4.5))
ax.plot(depths, depth_f1, marker="o", color="#C44E52")
ax.set_title("DecisionTree max_depth별 F1 변화")
ax.set_xlabel("max_depth"); ax.set_ylabel("F1 score")
ax.set_xticks(depths); ax.grid(True, alpha=0.3)
fig.savefig(os.path.join(IMG, "chart5_line_depth.png"), bbox_inches="tight", dpi=120)
plt.close(fig)

# ---------- 보고서(docx) 작성 ----------
doc = Document()
doc.add_heading("청렴지수·부패신고 데이터 분석 보고서", level=0)
p = doc.add_paragraph("[연습세트01·블루] 2과목 — 데이터분석")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("1. 분석 개요", level=1)
doc.add_paragraph(
    "행정안전부 청렴감사관실 업무로, 기관별 청렴지수 현황과 부패신고 현황 데이터를 "
    "병합·전처리한 뒤 머신러닝 모델로 기관의 부패위험 여부를 예측하였다.")

doc.add_heading("2. 데이터 전처리", level=1)
doc.add_paragraph(f"· 기관코드 기준 inner merge: 175행", style="List Bullet")
doc.add_paragraph(f"· 결측치(dropna) 제거 후: 148행", style="List Bullet")
doc.add_paragraph(f"· IQR 이상치(설문인원·평균처리일수) 제거 후: {len(clean)}행", style="List Bullet")
doc.add_paragraph(f"· 라벨(부패위험, seed=251) 분포 — 0: {(y==0).sum()}건 / 1: {(y==1).sum()}건",
                  style="List Bullet")

doc.add_heading("3. 핵심 분석 결과 (단답)", level=1)
g = clean.groupby("기관유형")["청렴지수"].mean().sort_values(ascending=False)
mean1 = clean.loc[clean["부패위험"] == 1, "청렴지수"].mean()
results = [
    ("Q1. 병합 기준 공통 컬럼", "기관코드"),
    ("Q2. 평균 청렴지수 최고 기관유형", f"{g.index[0]} ({g.iloc[0]:.3f})"),
    ("Q3. 정제 후 행 수", f"{len(clean)} 행"),
    ("Q4. 라벨=1 기관 청렴지수 평균(반올림)", f"{round(mean1)} (원값 {mean1:.4f})"),
    ("Q5. 테스트셋 라벨=1 행 수", f"{int((yte==1).sum())} 행 (테스트 {len(yte)}행)"),
]
t = doc.add_table(rows=1, cols=2)
t.style = "Light Grid Accent 1"
t.rows[0].cells[0].text = "항목"; t.rows[0].cells[1].text = "값"
for k, v in results:
    row = t.add_row().cells
    row[0].text = k; row[1].text = str(v)

doc.add_heading("4. 필수 차트 5종", level=1)
charts = [
    ("4.1 기관유형별 청렴지수 분포 (boxplot)", "chart1_boxplot.png",
     "광역지자체의 청렴지수 중앙값이 가장 높고, 기초지자체가 가장 낮은 분포를 보인다."),
    ("4.2 청렴지수 히스토그램", "chart2_histogram.png",
     "청렴지수는 7~8점대에 집중 분포하며 대체로 단봉형 형태를 띤다."),
    ("4.3 신고건수 vs 처리율 산점도", "chart3_scatter.png",
     "신고건수·처리율과 부패위험(색상)의 관계를 시각화하였다."),
    ("4.4 4개 알고리즘 F1 비교 (bar)", "chart4_bar_f1.png",
     "DecisionTree·RandomForest·LogisticRegression·KNN의 테스트셋 F1을 비교하였다. "
     + " / ".join(f"{n}: {v:.3f}" for n, v in f1_scores.items())),
    ("4.5 DecisionTree max_depth별 F1 변화 (line)", "chart5_line_depth.png",
     f"max_depth를 1~15로 변화시키며 F1을 측정하였다. 최고 F1={max(depth_f1):.3f} "
     f"(max_depth={depths[int(np.argmax(depth_f1))]})."),
]
for title, fname, desc in charts:
    doc.add_heading(title, level=2)
    doc.add_picture(os.path.join(IMG, fname), width=Inches(5.8))
    doc.add_paragraph(desc)

doc.add_heading("5. 결론", level=1)
best = max(f1_scores, key=f1_scores.get)
doc.add_paragraph(
    f"전처리 후 {len(clean)}개 기관 데이터로 부패위험 예측 모델을 학습한 결과, "
    f"{best} 모델이 F1 {f1_scores[best]:.3f}로 가장 우수하였다. "
    "청렴지수가 낮고 신고건수가 많은 기관일수록 부패위험이 높게 예측되는 경향을 확인하였다.")

out = os.path.join(BASE, "분석_보고서.docx")
doc.save(out)
print("saved:", out)
print("F1:", {k: round(v, 3) for k, v in f1_scores.items()})
print("depth best:", depths[int(np.argmax(depth_f1))], round(max(depth_f1), 3))
