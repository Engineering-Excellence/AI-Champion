"""
엑셀 자동 분석 대시보드 + Word 보고서 생성 (Streamlit UI)
==========================================================

<prompt>
  <task>
    엑셀파일을 업로드하면 해당 컬럼의 내용을 전부 확인, 분석한 뒤 그래프 10개와 통계표 5개의
    대시보드를 출력하고 docx 파일의 보고서를 만들어주는 파이썬 코드를 작성하고, 설치할 라이브러리를 알려줘.
  </task>
  <constraints>
    - 실제 설치하는 코드는 주석 처리할 것
    - 파일 선택이 가능하고, 분석 항목을 선택하는 등 사용자 친화적인 UI를 추가해서 제작할 것
  </constraints>
</prompt>


[설치할 라이브러리] (실제 설치 코드는 모두 주석 처리되어 있음)
- streamlit     : 웹 UI
- pandas        : 데이터 처리
- numpy         : 수치 연산
- matplotlib    : 그래프 생성
- seaborn       : 시각화
- openpyxl      : xlsx 파일 읽기
- python-docx   : Word(docx) 보고서 생성

# pip install streamlit pandas numpy matplotlib seaborn openpyxl python-docx

실행 방법:
# streamlit run app.py
"""

import io
import warnings
warnings.filterwarnings("ignore")

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns
import streamlit as st

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

sns.set_style("whitegrid")
plt.rcParams["axes.unicode_minus"] = False
# 한글 폰트가 필요한 경우 환경에 맞게 설정
plt.rcParams["font.family"] = "Malgun Gothic"  # Windows
# plt.rcParams["font.family"] = "AppleGothic"    # macOS

st.set_page_config(page_title="엑셀 데이터 분석 대시보드", layout="wide")


# ----------------------------------------------------------------------
# 데이터 로드 및 컬럼 분류
# ----------------------------------------------------------------------
@st.cache_data
def load_data(file):
    df = pd.read_excel(file)
    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    categorical_cols = df.select_dtypes(include=["object", "category"]).columns.tolist()
    datetime_cols = df.select_dtypes(include=["datetime64[ns]", "datetime64"]).columns.tolist()

    for col in df.columns:
        if col not in datetime_cols and col not in numeric_cols:
            try:
                converted = pd.to_datetime(df[col], errors="coerce")
                if converted.notna().mean() > 0.8:
                    df[col] = converted
                    datetime_cols.append(col)
                    if col in categorical_cols:
                        categorical_cols.remove(col)
            except Exception:
                pass

    return df, numeric_cols, categorical_cols, datetime_cols


# ----------------------------------------------------------------------
# 통계표 생성 함수들
# ----------------------------------------------------------------------
def stat_describe(df, numeric_cols):
    if not numeric_cols:
        return None
    return df[numeric_cols].describe().round(2)


def stat_missing(df):
    missing = df.isnull().sum()
    missing_pct = (df.isnull().mean() * 100).round(2)
    return pd.DataFrame({"결측치 수": missing, "결측 비율(%)": missing_pct})


def stat_correlation(df, numeric_cols):
    if len(numeric_cols) < 2:
        return None
    return df[numeric_cols].corr().round(2)


def stat_category_freq(df, cat_col):
    if not cat_col:
        return None
    return df[cat_col].value_counts().head(10).to_frame(name="빈도수")


def stat_distribution(df, numeric_cols):
    if not numeric_cols:
        return None
    return pd.DataFrame({
        "분산": df[numeric_cols].var().round(2),
        "표준편차": df[numeric_cols].std().round(2),
        "왜도": df[numeric_cols].skew().round(2),
        "첨도": df[numeric_cols].kurt().round(2),
    })


STAT_TABLE_REGISTRY = {
    "기술통계 요약": lambda df, num, cat, dt, target_cat: stat_describe(df, num),
    "결측치 현황": lambda df, num, cat, dt, target_cat: stat_missing(df),
    "상관관계 행렬": lambda df, num, cat, dt, target_cat: stat_correlation(df, num),
    "범주형 컬럼 빈도 Top10": lambda df, num, cat, dt, target_cat: stat_category_freq(df, target_cat),
    "분포 특성 (분산/왜도/첨도)": lambda df, num, cat, dt, target_cat: stat_distribution(df, num),
}


# ----------------------------------------------------------------------
# 그래프 생성 함수들 (각 함수는 fig를 반환)
# ----------------------------------------------------------------------
def chart_histogram(df, numeric_cols, **kwargs):
    cols = numeric_cols[:4]
    if not cols:
        return None
    fig, axes = plt.subplots(1, len(cols), figsize=(5 * len(cols), 4))
    axes = np.atleast_1d(axes)
    for ax, col in zip(axes, cols):
        sns.histplot(df[col].dropna(), kde=True, ax=ax)
        ax.set_title(f"{col} 분포")
    fig.tight_layout()
    return fig


def chart_boxplot(df, numeric_cols, **kwargs):
    if not numeric_cols:
        return None
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.boxplot(data=df[numeric_cols], ax=ax)
    ax.set_title("수치형 컬럼 박스플롯")
    ax.tick_params(axis="x", rotation=45)
    fig.tight_layout()
    return fig


def chart_corr_heatmap(df, numeric_cols, **kwargs):
    if len(numeric_cols) < 2:
        return None
    fig, ax = plt.subplots(figsize=(8, 6))
    sns.heatmap(df[numeric_cols].corr(), annot=True, cmap="coolwarm", fmt=".2f", ax=ax)
    ax.set_title("상관관계 히트맵")
    fig.tight_layout()
    return fig


def chart_category_bar(df, numeric_cols, target_cat=None, **kwargs):
    if not target_cat:
        return None
    vc = df[target_cat].value_counts().head(10)
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.barplot(x=vc.values, y=vc.index, ax=ax, palette="Blues_r")
    ax.set_title(f"{target_cat} - 빈도 Top10")
    fig.tight_layout()
    return fig


def chart_category_pie(df, numeric_cols, target_cat=None, **kwargs):
    if not target_cat:
        return None
    vc = df[target_cat].value_counts().head(6)
    fig, ax = plt.subplots(figsize=(6, 6))
    ax.pie(vc.values, labels=vc.index, autopct="%1.1f%%", startangle=90)
    ax.set_title(f"{target_cat} - 비율")
    fig.tight_layout()
    return fig


def chart_scatter(df, numeric_cols, target_x=None, target_y=None, **kwargs):
    x = target_x or (numeric_cols[0] if len(numeric_cols) > 0 else None)
    y = target_y or (numeric_cols[1] if len(numeric_cols) > 1 else None)
    if not x or not y or x == y:
        return None
    fig, ax = plt.subplots(figsize=(7, 5))
    sns.scatterplot(x=df[x], y=df[y], ax=ax)
    ax.set_title(f"{x} vs {y} 산점도")
    fig.tight_layout()
    return fig


def chart_timeseries(df, numeric_cols, datetime_cols=None, target_y=None, **kwargs):
    if not datetime_cols:
        return None
    date_col = datetime_cols[0]
    val_col = target_y or (numeric_cols[0] if numeric_cols else None)
    if not val_col:
        return None
    ts = df[[date_col, val_col]].dropna().sort_values(date_col)
    if ts.empty:
        return None
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.plot(ts[date_col], ts[val_col])
    ax.set_title(f"{val_col} 추이 ({date_col} 기준)")
    ax.tick_params(axis="x", rotation=45)
    fig.tight_layout()
    return fig


def chart_groupby_bar(df, numeric_cols, target_cat=None, target_y=None, **kwargs):
    if not target_cat:
        return None
    num_col = target_y or (numeric_cols[0] if numeric_cols else None)
    if not num_col:
        return None
    grouped = df.groupby(target_cat)[num_col].mean().sort_values(ascending=False).head(10)
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.barplot(x=grouped.index, y=grouped.values, ax=ax, palette="viridis")
    ax.set_title(f"{target_cat}별 {num_col} 평균")
    ax.tick_params(axis="x", rotation=45)
    fig.tight_layout()
    return fig


def chart_pairplot(df, numeric_cols, **kwargs):
    cols = numeric_cols[:3]
    if len(cols) < 2:
        return None
    pair = sns.pairplot(df[cols].dropna())
    pair.fig.suptitle("수치형 컬럼 페어플롯", y=1.02)
    return pair.fig


def chart_missing(df, numeric_cols, **kwargs):
    missing = df.isnull().sum()
    missing = missing[missing > 0]
    fig, ax = plt.subplots(figsize=(8, 5))
    if len(missing) > 0:
        sns.barplot(x=missing.values, y=missing.index, ax=ax, palette="Reds_r")
        ax.set_title("컬럼별 결측치 수")
    else:
        ax.text(0.5, 0.5, "결측치 없음", ha="center", va="center", fontsize=14)
        ax.set_title("결측치 현황")
        ax.axis("off")
    fig.tight_layout()
    return fig


def chart_violin(df, numeric_cols, **kwargs):
    if not numeric_cols:
        return None
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.violinplot(data=df[numeric_cols[:4]], ax=ax)
    ax.set_title("수치형 컬럼 분포 (바이올린플롯)")
    ax.tick_params(axis="x", rotation=45)
    fig.tight_layout()
    return fig


CHART_REGISTRY = {
    "히스토그램 (분포)": chart_histogram,
    "박스플롯": chart_boxplot,
    "상관관계 히트맵": chart_corr_heatmap,
    "범주형 빈도 막대그래프": chart_category_bar,
    "범주형 비율 파이차트": chart_category_pie,
    "산점도": chart_scatter,
    "시계열 라인그래프": chart_timeseries,
    "범주형별 평균 막대그래프": chart_groupby_bar,
    "페어플롯": chart_pairplot,
    "바이올린플롯": chart_violin,
    "결측치 현황": chart_missing,
}


# ----------------------------------------------------------------------
# Word 보고서 생성
# ----------------------------------------------------------------------
def dataframe_to_docx_table(doc, df, max_rows=15):
    df_show = df.reset_index().head(max_rows)
    table = doc.add_table(rows=1, cols=len(df_show.columns))
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_show.columns):
        hdr_cells[i].text = str(col)

    for _, row in df_show.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def build_report(file_name, df, numeric_cols, categorical_cols, datetime_cols,
                  selected_tables, selected_charts, chart_images):
    doc = Document()

    title = doc.add_heading("데이터 분석 보고서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"분석 대상 파일: {file_name}")
    doc.add_paragraph(f"전체 행 수: {len(df)}")
    doc.add_paragraph(f"전체 컬럼 수: {len(df.columns)}")
    doc.add_paragraph(f"수치형 컬럼: {', '.join(numeric_cols) if numeric_cols else '없음'}")
    doc.add_paragraph(f"범주형 컬럼: {', '.join(categorical_cols) if categorical_cols else '없음'}")
    doc.add_paragraph(f"날짜형 컬럼: {', '.join(datetime_cols) if datetime_cols else '없음'}")

    doc.add_page_break()

    # 통계표
    doc.add_heading("1. 통계표", level=1)
    for name, table_df in selected_tables.items():
        if table_df is None:
            continue
        doc.add_heading(name, level=2)
        dataframe_to_docx_table(doc, table_df)
        doc.add_paragraph()

    doc.add_page_break()

    # 그래프
    doc.add_heading("2. 시각화 대시보드", level=1)
    for name, img_bytes in chart_images.items():
        if img_bytes is None:
            continue
        doc.add_heading(name, level=2)
        doc.add_picture(io.BytesIO(img_bytes), width=Inches(6))
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ----------------------------------------------------------------------
# Streamlit UI
# ----------------------------------------------------------------------
def main():
    st.title("📊 엑셀 데이터 분석 대시보드")
    st.write("엑셀 파일을 업로드하고, 분석할 그래프와 통계표를 선택해 대시보드와 Word 보고서를 생성합니다.")

    uploaded_file = st.file_uploader("엑셀 파일 선택 (.xlsx)", type=["xlsx", "xls"])

    if uploaded_file is None:
        st.info("엑셀 파일을 업로드해주세요.")
        return

    df, numeric_cols, categorical_cols, datetime_cols = load_data(uploaded_file)

    # 데이터 미리보기
    with st.expander("데이터 미리보기", expanded=True):
        st.dataframe(df.head(20))
        col1, col2, col3 = st.columns(3)
        col1.metric("전체 행 수", len(df))
        col2.metric("전체 컬럼 수", len(df.columns))
        col3.metric("수치형 컬럼 수", len(numeric_cols))

    st.sidebar.header("⚙️ 분석 옵션")

    # 범주형/수치형 컬럼 선택 (그래프 생성에 사용)
    target_cat = None
    if categorical_cols:
        target_cat = st.sidebar.selectbox("범주형 분석 컬럼", ["(자동 선택)"] + categorical_cols)
        if target_cat == "(자동 선택)":
            target_cat = categorical_cols[0]

    target_x, target_y = None, None
    if len(numeric_cols) >= 2:
        target_x = st.sidebar.selectbox("산점도 X축", numeric_cols, index=0)
        target_y = st.sidebar.selectbox("산점도 / 그룹분석 Y축", numeric_cols, index=min(1, len(numeric_cols) - 1))

    st.sidebar.subheader("📈 표시할 그래프 선택")
    selected_chart_names = []
    for name in CHART_REGISTRY.keys():
        if st.sidebar.checkbox(name, value=True, key=f"chart_{name}"):
            selected_chart_names.append(name)

    st.sidebar.subheader("📋 표시할 통계표 선택")
    selected_table_names = []
    for name in STAT_TABLE_REGISTRY.keys():
        if st.sidebar.checkbox(name, value=True, key=f"table_{name}"):
            selected_table_names.append(name)

    run = st.button("🚀 분석 실행", type="primary")

    if not run:
        return

    if not selected_chart_names and not selected_table_names:
        st.warning("최소 하나 이상의 그래프 또는 통계표를 선택해주세요.")
        return

    # ----------------------------------------------------------------------
    # 통계표 출력
    # ----------------------------------------------------------------------
    selected_tables = {}
    if selected_table_names:
        st.header("📋 통계표")
        for name in selected_table_names:
            result_df = STAT_TABLE_REGISTRY[name](df, numeric_cols, categorical_cols, datetime_cols, target_cat)
            selected_tables[name] = result_df
            st.subheader(name)
            if result_df is not None:
                st.dataframe(result_df)
            else:
                st.caption("이 표를 생성하기 위한 적절한 컬럼이 없습니다.")

    # ----------------------------------------------------------------------
    # 그래프 출력
    # ----------------------------------------------------------------------
    chart_images = {}
    if selected_chart_names:
        st.header("📈 그래프 대시보드")
        cols_per_row = 2
        rows = [selected_chart_names[i:i + cols_per_row] for i in range(0, len(selected_chart_names), cols_per_row)]

        for row_names in rows:
            cols = st.columns(len(row_names))
            for col, name in zip(cols, row_names):
                fig = CHART_REGISTRY[name](
                    df, numeric_cols,
                    target_cat=target_cat,
                    target_x=target_x,
                    target_y=target_y,
                    datetime_cols=datetime_cols,
                )
                with col:
                    st.subheader(name)
                    if fig is not None:
                        st.pyplot(fig)
                        chart_images[name] = fig_to_bytes(fig)
                    else:
                        st.caption("이 그래프를 생성하기 위한 적절한 컬럼이 없습니다.")
                        chart_images[name] = None

    # ----------------------------------------------------------------------
    # Word 보고서 다운로드
    # ----------------------------------------------------------------------
    st.header("📄 보고서 다운로드")
    report_buf = build_report(
        uploaded_file.name, df, numeric_cols, categorical_cols, datetime_cols,
        selected_tables, selected_chart_names, chart_images
    )
    st.download_button(
        label="📥 Word 보고서 다운로드 (.docx)",
        data=report_buf,
        file_name="분석_보고서.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    main()