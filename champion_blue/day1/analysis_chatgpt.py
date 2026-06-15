"""
<prompt>
  <task>
    엑셀파일을 업로드하면 해당 컬럼의 내용을 전부 확인, 분석한 뒤 그래프 10개와 통계표 5개의
    대시보드를 출력하고 docx 파일의 보고서를 만들어주는 파이썬 코드를 작성하고, 설치할 라이브러리를 알려줘.
  </task>
  <constraints>
    - 실제 설치하는 코드는 주석 처리할 것
  </constraints>
</prompt>
"""

import os
import warnings

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.shared import Inches

warnings.filterwarnings("ignore")

# ---------------------------
# 설정
# ---------------------------

INPUT_FILE = "input/public_sector_analysis_dataset_3000x20.xlsx"
OUTPUT_DIR = "output"

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ---------------------------
# 데이터 읽기
# ---------------------------

df = pd.read_excel(INPUT_FILE)

print("=" * 60)
print("데이터 크기")
print(df.shape)
print("=" * 60)

# ---------------------------
# 컬럼 분류
# ---------------------------

numeric_cols = df.select_dtypes(
    include=["int64", "float64"]
).columns.tolist()

datetime_cols = df.select_dtypes(
    include=["datetime64"]
).columns.tolist()

categorical_cols = [
    c for c in df.columns
    if c not in numeric_cols + datetime_cols
]

# ---------------------------
# 통계표 생성
# ---------------------------

tables = {}

# 1. 기초통계
tables["기초통계"] = df.describe(include="all")

# 2. 결측치
tables["결측치"] = pd.DataFrame({
    "컬럼": df.columns,
    "결측치수": df.isna().sum(),
    "결측비율(%)": round(
        df.isna().mean() * 100,
        2
    )
})

# 3. 데이터형
tables["데이터형"] = pd.DataFrame({
    "컬럼": df.columns,
    "dtype": df.dtypes.astype(str)
})

# 4. 상관계수
if len(numeric_cols) >= 2:
    tables["상관계수"] = df[numeric_cols].corr()
else:
    tables["상관계수"] = pd.DataFrame()

# 5. 이상치
outlier_result = []

for col in numeric_cols:

    q1 = df[col].quantile(0.25)
    q3 = df[col].quantile(0.75)

    iqr = q3 - q1

    lower = q1 - 1.5 * iqr
    upper = q3 + 1.5 * iqr

    count = (
        (df[col] < lower) |
        (df[col] > upper)
    ).sum()

    outlier_result.append([
        col,
        count
    ])

tables["이상치"] = pd.DataFrame(
    outlier_result,
    columns=["컬럼", "이상치수"]
)

# 엑셀 저장
stats_file = os.path.join(
    OUTPUT_DIR,
    "statistics.xlsx"
)

with pd.ExcelWriter(stats_file) as writer:

    for name, table in tables.items():
        table.to_excel(
            writer,
            sheet_name=name[:30]
        )

print("통계표 저장 완료")

# ---------------------------
# 그래프 생성
# ---------------------------

chart_files = []

chart_no = 1

def save_chart():
    global chart_no

    filename = os.path.join(
        OUTPUT_DIR,
        f"chart_{chart_no:02d}.png"
    )

    plt.tight_layout()
    plt.savefig(
        filename,
        dpi=300,
        bbox_inches="tight"
    )
    plt.close()

    chart_files.append(filename)

    chart_no += 1

# ---------------------------
# 1~3 히스토그램
# ---------------------------

for col in numeric_cols[:3]:

    plt.figure(figsize=(8,5))

    sns.histplot(
        df[col].dropna(),
        kde=True
    )

    plt.title(f"Histogram - {col}")

    save_chart()

# ---------------------------
# 4~5 박스플롯
# ---------------------------

for col in numeric_cols[:2]:

    plt.figure(figsize=(8,5))

    sns.boxplot(
        x=df[col]
    )

    plt.title(f"Boxplot - {col}")

    save_chart()

# ---------------------------
# 6 상관관계 히트맵
# ---------------------------

if len(numeric_cols) >= 2:

    plt.figure(figsize=(10,8))

    sns.heatmap(
        df[numeric_cols].corr(),
        annot=True,
        cmap="coolwarm"
    )

    plt.title("Correlation Heatmap")

    save_chart()

# ---------------------------
# 7~8 범주형 빈도
# ---------------------------

for col in categorical_cols[:2]:

    plt.figure(figsize=(8,5))

    df[col].value_counts().head(10).plot(
        kind="bar"
    )

    plt.title(f"Top Categories - {col}")

    save_chart()

# ---------------------------
# 9 산점도
# ---------------------------

if len(numeric_cols) >= 2:

    plt.figure(figsize=(8,5))

    sns.scatterplot(
        data=df,
        x=numeric_cols[0],
        y=numeric_cols[1]
    )

    plt.title("Scatter Plot")

    save_chart()

# ---------------------------
# 10 Pairplot
# ---------------------------

if len(numeric_cols) >= 2:

    sns.pairplot(
        df[numeric_cols[:4]]
    )

    filename = os.path.join(
        OUTPUT_DIR,
        f"chart_{chart_no:02d}.png"
    )

    plt.savefig(
        filename,
        dpi=300
    )

    plt.close()

    chart_files.append(filename)

# ---------------------------
# DOCX 보고서 생성
# ---------------------------

doc = Document()

doc.add_heading(
    "자동 데이터 분석 보고서",
    level=1
)

doc.add_paragraph(
    f"총 행 수: {len(df):,}"
)

doc.add_paragraph(
    f"총 컬럼 수: {len(df.columns):,}"
)

doc.add_paragraph(
    f"숫자형 컬럼: {len(numeric_cols)}"
)

doc.add_paragraph(
    f"범주형 컬럼: {len(categorical_cols)}"
)

doc.add_paragraph(
    f"날짜형 컬럼: {len(datetime_cols)}"
)

# 통계표 요약

for name, table in tables.items():

    doc.add_heading(
        name,
        level=2
    )

    doc.add_paragraph(
        table.head(15).to_string()
    )

# 그래프 삽입

doc.add_heading(
    "그래프",
    level=2
)

for img in chart_files:

    doc.add_picture(
        img,
        width=Inches(5.5)
    )

# 저장

doc_file = os.path.join(
    OUTPUT_DIR,
    "analysis.docx"
)

doc.save(doc_file)

print("DOCX 생성 완료")
print(doc_file)
